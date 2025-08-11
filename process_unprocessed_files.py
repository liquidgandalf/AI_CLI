#!/usr/bin/env python3
"""
Background processor for unprocessed ChatFile attachments.

This script processes uploaded files based on their type:
- Audio files (mp3, wav): Transcribe using faster-whisper
- Text files: Extract content (future)
- PDF files: Extract text (future)
- Image files: OCR/description (future)

Usage:
    python process_unprocessed_files.py --once    # Process one batch and exit
    python process_unprocessed_files.py --loop    # Run continuously (daemon mode)
"""

import sys
import os
import time
import argparse
from datetime import datetime
from typing import Optional, Tuple

# Add the current directory to Python path to import app modules
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from database import ChatFile, get_db
from file_utils import get_file_full_path

# Import faster-whisper for audio transcription
try:
    from faster_whisper import WhisperModel
    WHISPER_AVAILABLE = True
except ImportError:
    WHISPER_AVAILABLE = False
    print("⚠️  faster-whisper not available. Audio transcription will be skipped.")

# Import PyMuPDF for PDF text extraction
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️  PyMuPDF not available. PDF text extraction will be skipped.")

# Import pandas for Excel file processing
try:
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("⚠️  pandas not available. Excel file processing will be skipped.")

# Import libraries for Word document processing
try:
    from docx import Document
    import docx2txt
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False
    print("⚠️  python-docx/docx2txt not available. Word document processing will be skipped.")

# Import libraries for LibreOffice/OpenOffice document processing
try:
    from odf.opendocument import load
    from odf.text import P
    from odf.table import Table, TableRow, TableCell
    ODF_AVAILABLE = True
except ImportError:
    ODF_AVAILABLE = False
    print("⚠️  odfpy not available. LibreOffice/OpenOffice document processing will be skipped.")

class FileProcessor:
    """Main file processing engine"""
    
    def __init__(self):
        self.whisper_model = None
        self._init_whisper()
    
    def _init_whisper(self):
        """Initialize Whisper model for audio transcription"""
        if not WHISPER_AVAILABLE:
            return
        
        try:
            print("🔄 Loading Whisper model (this may take a moment on first run)...")
            # Use base model for good balance of speed and accuracy
            # Options: tiny, base, small, medium, large-v2, large-v3
            self.whisper_model = WhisperModel("base", device="cpu", compute_type="int8")
            print("✅ Whisper model loaded successfully")
        except Exception as e:
            print(f"❌ Failed to load Whisper model: {e}")
            self.whisper_model = None
    
    def process_audio_file(self, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """
        Transcribe audio file using faster-whisper
        
        Returns:
            (success, transcription_text, error_message)
        """
        if not self.whisper_model:
            return False, "", "Whisper model not available"
        
        if not os.path.exists(file_path):
            return False, "", f"File not found: {file_path}"
        
        try:
            print(f"🎵 Transcribing audio file: {os.path.basename(file_path)}")
            
            # Transcribe the audio file
            segments, info = self.whisper_model.transcribe(file_path, beam_size=5)
            
            # Collect all segments into a single transcript
            transcript_parts = []
            for segment in segments:
                transcript_parts.append(f"[{segment.start:.1f}s - {segment.end:.1f}s] {segment.text}")
            
            full_transcript = "\n".join(transcript_parts)
            
            # Also create a clean version without timestamps
            clean_transcript = "\n".join([segment.text.strip() for segment in segments if segment.text.strip()])
            
            # Combine both versions
            final_transcript = f"=== CLEAN TRANSCRIPT ===\n{clean_transcript}\n\n=== DETAILED TRANSCRIPT WITH TIMESTAMPS ===\n{full_transcript}"
            
            print(f"✅ Transcription completed. Language: {info.language} (confidence: {info.language_probability:.2f})")
            return True, final_transcript, None
            
        except Exception as e:
            error_msg = f"Transcription failed: {str(e)}"
            print(f"❌ {error_msg}")
            return False, "", error_msg
    
    def process_text_file(self, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """
        Extract text content from text files
        
        Returns:
            (success, text_content, error_message)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return True, content, None
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                return True, content, None
            except Exception as e:
                return False, "", f"Failed to read text file: {str(e)}"
        except Exception as e:
            return False, "", f"Failed to read text file: {str(e)}"
    
    def process_pdf_file(self, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """
        Extract text content from PDF files using PyMuPDF
        
        Returns:
            (success, text_content, error_message)
        """
        if not PDF_AVAILABLE:
            return False, "", "PyMuPDF not available for PDF processing"
        
        if not os.path.exists(file_path):
            return False, "", f"File not found: {file_path}"
        
        try:
            print(f"📄 Extracting text from PDF: {os.path.basename(file_path)}")
            
            # Open the PDF document
            doc = fitz.open(file_path)
            
            if doc.is_encrypted:
                return False, "", "PDF is encrypted and cannot be processed"
            
            # Extract text from all pages
            text_content = []
            page_count = len(doc)
            
            for page_num in range(page_count):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                
                if page_text.strip():  # Only add non-empty pages
                    text_content.append(f"=== PAGE {page_num + 1} ===\n{page_text.strip()}")
            
            doc.close()
            
            if not text_content:
                return False, "", "No text content found in PDF (may be image-based)"
            
            # Combine all pages
            full_text = "\n\n".join(text_content)
            
            # Add metadata summary
            summary = f"=== PDF METADATA ===\nTotal Pages: {page_count}\nPages with Text: {len(text_content)}\nFile: {os.path.basename(file_path)}\n\n=== EXTRACTED TEXT ===\n{full_text}"
            
            print(f"✅ PDF text extraction completed. Pages: {page_count}, Text pages: {len(text_content)}")
            return True, summary, None
            
        except Exception as e:
            error_msg = f"PDF processing failed: {str(e)}"
            print(f"❌ {error_msg}")
            return False, "", error_msg
    
    def process_excel_file(self, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """
        Extract content from Excel files using pandas
        
        Returns:
            (success, text_content, error_message)
        """
        if not EXCEL_AVAILABLE:
            return False, "", "pandas not available for Excel processing"
        
        if not os.path.exists(file_path):
            return False, "", f"File not found: {file_path}"
        
        try:
            print(f"📊 Processing Excel file: {os.path.basename(file_path)}")
            
            # Read all sheets from the Excel file
            excel_file = pd.ExcelFile(file_path)
            all_content = []
            
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Convert DataFrame to string representation
                    sheet_content = f"=== SHEET: {sheet_name} ==="
                    sheet_content += f"\nRows: {len(df)}, Columns: {len(df.columns)}\n"
                    sheet_content += f"\nColumn Names: {', '.join(df.columns.astype(str))}\n\n"
                    
                    # Add the actual data (limit to first 1050 rows to avoid huge content)
                    if len(df) > 0:
                        display_df = df.head(1050)  # Show first 1050 rows
                        sheet_content += display_df.to_string(index=False, max_rows=1050)
                        
                        if len(df) > 1050:
                            sheet_content += f"\n\n... ({len(df) - 1050} more rows not shown)"
                    else:
                        sheet_content += "(Empty sheet)"
                    
                    all_content.append(sheet_content)
                    
                except Exception as sheet_error:
                    all_content.append(f"=== SHEET: {sheet_name} ===\nError reading sheet: {sheet_error}")
            
            final_content = "\n\n".join(all_content)
            print(f"✅ Excel processing completed. Found {len(excel_file.sheet_names)} sheets")
            return True, final_content, None
            
        except Exception as e:
            error_msg = f"Excel processing failed: {str(e)}"
            print(f"❌ {error_msg}")
            return False, "", error_msg
    
    def process_word_file(self, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """
        Extract content from Word documents (.docx and .doc)
        
        Returns:
            (success, text_content, error_message)
        """
        if not WORD_AVAILABLE:
            return False, "", "python-docx/docx2txt not available for Word processing"
        
        if not os.path.exists(file_path):
            return False, "", f"File not found: {file_path}"
        
        try:
            print(f"📄 Processing Word document: {os.path.basename(file_path)}")
            
            # Check file extension to determine processing method
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.docx':
                # Use python-docx for .docx files
                doc = Document(file_path)
                
                # Extract text from paragraphs
                paragraphs = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text.strip())
                
                # Extract text from tables
                tables_content = []
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        if any(row_data):  # Only add non-empty rows
                            table_data.append(" | ".join(row_data))
                    
                    if table_data:
                        tables_content.append("=== TABLE ===\n" + "\n".join(table_data))
                
                # Combine all content
                all_content = []
                if paragraphs:
                    all_content.append("=== DOCUMENT TEXT ===\n" + "\n\n".join(paragraphs))
                if tables_content:
                    all_content.append("\n\n".join(tables_content))
                
                final_content = "\n\n".join(all_content)
                
            elif file_ext == '.doc':
                # Use docx2txt for .doc files (simpler extraction)
                final_content = docx2txt.process(file_path)
                if final_content:
                    final_content = "=== DOCUMENT TEXT ===\n" + final_content.strip()
                else:
                    final_content = "(No readable text found in document)"
            
            else:
                return False, "", f"Unsupported Word document format: {file_ext}"
            
            if not final_content or final_content.strip() == "=== DOCUMENT TEXT ===" or final_content.strip() == "(No readable text found in document)":
                final_content = "(Document appears to be empty or contains no readable text)"
            
            print(f"✅ Word document processing completed. Extracted {len(final_content)} characters")
            return True, final_content, None
            
        except Exception as e:
            error_msg = f"Word document processing failed: {str(e)}"
            print(f"❌ {error_msg}")
            return False, "", error_msg
    
    def process_odf_file(self, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """
        Extract content from LibreOffice/OpenOffice documents (.odt, .ods)
        
        Returns:
            (success, text_content, error_message)
        """
        if not ODF_AVAILABLE:
            return False, "", "odfpy not available for LibreOffice/OpenOffice processing"
        
        if not os.path.exists(file_path):
            return False, "", f"File not found: {file_path}"
        
        try:
            print(f"📄 Processing LibreOffice/OpenOffice document: {os.path.basename(file_path)}")
            
            # Load the ODF document
            doc = load(file_path)
            file_ext = os.path.splitext(file_path)[1].lower()
            
            all_content = []
            
            if file_ext == '.odt':  # Writer document
                # Extract paragraphs
                paragraphs = []
                for paragraph in doc.getElementsByType(P):
                    text = str(paragraph).strip()
                    if text and not text.startswith('<') and not text.endswith('>'):
                        # Clean up the text (remove XML-like content)
                        clean_text = paragraph.firstChild.data if paragraph.firstChild else ""
                        if clean_text and clean_text.strip():
                            paragraphs.append(clean_text.strip())
                
                # Extract tables
                tables_content = []
                for table in doc.getElementsByType(Table):
                    table_data = []
                    for row in table.getElementsByType(TableRow):
                        row_data = []
                        for cell in row.getElementsByType(TableCell):
                            cell_text = ""
                            for p in cell.getElementsByType(P):
                                if p.firstChild:
                                    cell_text += str(p.firstChild.data or "").strip() + " "
                            row_data.append(cell_text.strip())
                        if any(row_data):  # Only add non-empty rows
                            table_data.append(" | ".join(row_data))
                    
                    if table_data:
                        tables_content.append("=== TABLE ===\n" + "\n".join(table_data))
                
                # Combine content
                if paragraphs:
                    all_content.append("=== DOCUMENT TEXT ===\n" + "\n\n".join(paragraphs))
                if tables_content:
                    all_content.extend(tables_content)
            
            elif file_ext == '.ods':  # Calc spreadsheet
                # Extract spreadsheet data
                sheets_content = []
                
                # Get all tables (sheets) in the spreadsheet
                tables = doc.getElementsByType(Table)
                
                for i, table in enumerate(tables):
                    sheet_name = table.getAttribute('name') or f"Sheet{i+1}"
                    sheet_data = []
                    
                    rows = table.getElementsByType(TableRow)
                    for row in rows:
                        row_data = []
                        cells = row.getElementsByType(TableCell)
                        for cell in cells:
                            cell_text = ""
                            for p in cell.getElementsByType(P):
                                if p.firstChild:
                                    cell_text += str(p.firstChild.data or "").strip() + " "
                            row_data.append(cell_text.strip())
                        
                        # Only add rows that have some content
                        if any(cell.strip() for cell in row_data):
                            sheet_data.append(" | ".join(row_data))
                    
                    if sheet_data:
                        sheet_content = f"=== SHEET: {sheet_name} ===\n"
                        sheet_content += f"Rows: {len(sheet_data)}\n\n"
                        sheet_content += "\n".join(sheet_data[:1050])  # Limit to first 1050 rows
                        
                        if len(sheet_data) > 1050:
                            sheet_content += f"\n\n... ({len(sheet_data) - 1050} more rows not shown)"
                        
                        sheets_content.append(sheet_content)
                
                all_content.extend(sheets_content)
            
            else:
                return False, "", f"Unsupported ODF document format: {file_ext}"
            
            if not all_content:
                final_content = "(Document appears to be empty or contains no readable content)"
            else:
                final_content = "\n\n".join(all_content)
            
            print(f"✅ LibreOffice/OpenOffice document processing completed. Extracted {len(final_content)} characters")
            return True, final_content, None
            
        except Exception as e:
            error_msg = f"LibreOffice/OpenOffice document processing failed: {str(e)}"
            print(f"❌ {error_msg}")
            return False, "", error_msg
    
    def process_file(self, chat_file: ChatFile) -> Tuple[bool, str, Optional[str]]:
        """
        Process a file based on its type
        
        Returns:
            (success, processed_content, error_message)
        """
        file_path = get_file_full_path(chat_file.file_path)
        
        if chat_file.file_type == 'audio':
            return self.process_audio_file(file_path)
        elif chat_file.file_type == 'text':
            return self.process_text_file(file_path)
        elif chat_file.file_type == 'document' and chat_file.mime_type == 'application/pdf':
            return self.process_pdf_file(file_path)
        elif chat_file.file_type == 'document' and chat_file.mime_type in ['application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
            return self.process_excel_file(file_path)
        elif chat_file.file_type == 'document' and chat_file.mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            return self.process_word_file(file_path)
        elif chat_file.file_type == 'document' and chat_file.mime_type in ['application/vnd.oasis.opendocument.text', 'application/vnd.oasis.opendocument.spreadsheet']:
            return self.process_odf_file(file_path)
        else:
            return False, "", f"File type '{chat_file.file_type}' with MIME type '{chat_file.mime_type}' not yet supported for processing"

def get_next_unprocessed_file(db) -> Optional[ChatFile]:
    """Get the next unprocessed file (status = 0) and mark it as being processed (status = 1)
    Skip files marked as "Do Not Process" (status = 4)
    """
    try:
        # Find the oldest unprocessed file
        chat_file = db.query(ChatFile).filter(
            ChatFile.has_been_processed == 0
        ).order_by(ChatFile.upload_date.asc()).first()
        
        if not chat_file:
            return None
        
        # Atomically mark it as being processed
        chat_file.has_been_processed = 1
        db.commit()
        
        return chat_file
        
    except Exception as e:
        print(f"❌ Error getting next unprocessed file: {e}")
        db.rollback()
        return None

def mark_file_processed(db, chat_file: ChatFile, transcoded_content: str, processing_time: float):
    """Mark file as successfully processed"""
    try:
        chat_file.transcoded_raw_file = transcoded_content
        chat_file.date_processed = datetime.utcnow()
        chat_file.time_to_process = processing_time
        chat_file.has_been_processed = 2
        db.commit()
        print(f"✅ File marked as processed: {chat_file.original_filename}")
    except Exception as e:
        print(f"❌ Error marking file as processed: {e}")
        db.rollback()

def mark_file_failed(db, chat_file: ChatFile, error_message: str, processing_time: float):
    """Mark file as failed processing"""
    try:
        chat_file.has_been_processed = 3  # Mark as permanently failed (not 0 which would retry)
        chat_file.human_notes = f"Processing failed at {datetime.utcnow()}: {error_message}"
        chat_file.time_to_process = processing_time
        db.commit()
        print(f"❌ File marked as permanently failed: {chat_file.original_filename}")
    except Exception as e:
        print(f"❌ Error marking file as failed: {e}")
        db.rollback()

def process_one_batch():
    """Process one batch of unprocessed files"""
    db = get_db()
    processor = FileProcessor()
    
    try:
        # Get next unprocessed file
        chat_file = get_next_unprocessed_file(db)
        
        if not chat_file:
            print("📭 No unprocessed files found")
            return False
        
        print(f"🔄 Processing file: {chat_file.original_filename} (ID: {chat_file.id}, Type: {chat_file.file_type})")
        
        # Process the file
        start_time = time.time()
        success, content, error = processor.process_file(chat_file)
        processing_time = time.time() - start_time
        
        if success:
            mark_file_processed(db, chat_file, content, processing_time)
            print(f"🎉 Successfully processed in {processing_time:.2f}s")
        else:
            mark_file_failed(db, chat_file, error or "Unknown error", processing_time)
            print(f"💥 Processing failed after {processing_time:.2f}s: {error}")
        
        return True
        
    except Exception as e:
        print(f"❌ Unexpected error in process_one_batch: {e}")
        return False
    finally:
        db.close()

def main():
    parser = argparse.ArgumentParser(description="Process unprocessed ChatFile attachments")
    parser.add_argument('--once', action='store_true', help='Process one batch and exit')
    parser.add_argument('--loop', action='store_true', help='Run continuously (daemon mode)')
    parser.add_argument('--sleep', type=int, default=10, help='Sleep seconds between batches in loop mode (default: 10)')
    
    args = parser.parse_args()
    
    if not args.once and not args.loop:
        print("❌ Please specify either --once or --loop")
        sys.exit(1)
    
    print("🚀 ChatFile Background Processor")
    print("=" * 40)
    
    if args.once:
        print("📋 Running in one-shot mode")
        success = process_one_batch()
        if success:
            print("✅ Batch processing completed")
        else:
            print("📭 No files to process")
    
    elif args.loop:
        print(f"🔄 Running in daemon mode (checking every {args.sleep}s)")
        print("Press Ctrl+C to stop")
        
        try:
            while True:
                had_work = process_one_batch()
                
                if had_work:
                    # If we processed something, check immediately for more
                    time.sleep(1)
                else:
                    # If no work, sleep longer
                    time.sleep(args.sleep)
                    
        except KeyboardInterrupt:
            print("\n🛑 Stopping background processor...")
            print("✅ Shutdown complete")

if __name__ == "__main__":
    main()
